"""Render docs/PAPER_REVISED.md as a Word document.

Handles only the subset of markdown used by that file: ATX headings,
pipe tables, fenced code blocks, bullet lists and **bold** spans.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

BOLD = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")


def add_runs(paragraph, text):
    """Split text on bold/italic markers so emphasis survives the conversion."""
    position = 0
    for match in BOLD.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        run = paragraph.add_run(match.group(1) or match.group(2))
        run.bold = match.group(1) is not None
        run.italic = match.group(2) is not None
        position = match.end()
    paragraph.add_run(text[position:])


def add_table(document, rows):
    header, *body = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    body = [r for r in body if not set("".join(r)) <= {"-", ":", " "}]

    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, header):
        add_runs(cell.paragraphs[0], f"**{text}**")
    for row in body:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            add_runs(cell.paragraphs[0], text)


def build(source: Path, target: Path):
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    lines = source.read_text().splitlines()
    index = 0

    while index < len(lines):
        line = lines[index]

        if line.startswith("```"):
            index += 1
            code = []
            while index < len(lines) and not lines[index].startswith("```"):
                code.append(lines[index])
                index += 1
            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)

        elif line.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                rows.append(lines[index])
                index += 1
            add_table(document, rows)
            index -= 1

        elif line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            heading = document.add_heading(level=min(level, 4))
            add_runs(heading, line.lstrip("# ").strip())
            if level == 1:
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif line.startswith("- "):
            add_runs(document.add_paragraph(style="List Bullet"), line[2:])

        elif re.match(r"^\d+\. ", line):
            add_runs(
                document.add_paragraph(style="List Number"),
                re.sub(r"^\d+\. ", "", line),
            )

        elif line.strip():
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(paragraph, line.strip())

        index += 1

    document.save(target)


if __name__ == "__main__":
    root = Path(__file__).parent
    name = sys.argv[1] if len(sys.argv) > 1 else "MedGuide_AI_Paper_Revised.docx"
    build(root / "PAPER_REVISED.md", root / name)
